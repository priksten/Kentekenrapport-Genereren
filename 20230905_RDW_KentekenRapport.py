from rdw.rdw import Rdw
import docx
import pandas as pd

# Instructie over programma: kenteken invoeren en afsluiten
instruction = "\nMet dit programma kunt u gegevens over voertuigen opvragen door een kenteken in te voeren"
instruction += "\nEnter 'q' om het programma af te sluiten."
instruction += "\nKenteken: "

active = True

while active:
    kenteken = input(instruction)

    if kenteken == 'q':
        active = False
    else:
        print('Uw kenteken: ' + kenteken)
        # print(type(kenteken))

        # Haal de voertuiggegevens bij het RDW op 
        car = Rdw()
        result = car.get_vehicle_data(kenteken)                        

        # print(result)

        # We gaan de voertuiggegevens exporteren naar Word. De bedoeling is dat dit in tabelvorm in een Word-bestand komt.
        # Hiervoor zetten we de volgende stappen:
        #   1. We zetten de voertuig-informatie van het RDW in een dataframe.We gebruiken dit DataFrame om de opgevraagde gegevens overzichteljik in de terminal te tonen
        #   2. We maken in Word een lege tabel
        #   3. We plaatsen de gegevens uit het dataframe in de lege tabel in Word
        #   4. We slaan het Word-document op onder de naam 'kenteken.docx'
        #   5. Tot slot staat er op mijn verlanglijstje om de tabel mooi op te maken, zodat de informatie overzichtelijker wordt weergegeven.

        result_str = result[0]
        
        # stap 1: gegevens in dataframe
        car_property = []
        car_data = []

        for key in result_str.keys():
            car_property.append(key)
        
        for value in result_str.values():
            car_data.append(value)

        # We maken nu woordenboeken die als basis kunnen dienen voor filteropties
        # In het woordenboek 'car_information' staan alle kentekengegevens die je via de RDW API ophaalt uit de database
        car_information = {
                'Property' : car_property,
                'Car Data' : car_data
            }

        # We filteren op basis van categorieen:
        #       1. Basiskenmerken
        #       2. Registratie
        #       3. Motor
        #       4. Milieu
        #       5. Maten en gewichten                 
      
        list_basiskenmerken_keys = ['voertuigsoort', 'merk', 'handelsbenaming', 'vervaldatum_apk', 'datum_tenaamstelling', 'bruto_bpm', 'inrichting', 'aantal_zitplaatsen', 
                                    'eerste_kleur', 'tweede_kleur', 'catalogusprijs', 'aantal_deuren', 'aantal_wielen', 'type', 'typegoedkeuringsnummer', 
                                    'variant', 'uitvoering']
        list_registratie_keys = ['vervaldatum_apk', 'datum_tenaamstelling', 'datum_eerste_toelating', 'datum_eerste_tenaamstelling_in_nederland', 'wam_verzekerd', 
                                 'plaats_chassisnummer', 'jaar_laatste_registratie_tellerstand', 'tellerstandoordeel', 'tenaamstellen_mogelijk', 'vervaldatum_apk_dt',
                                 'datum_tenaamstelling_dt', 'datum_eerste_toelating_dt', 'datum_eerste_tenaamstelling_in_nederland_dt']
        
        list_motor_keys = ['aantal_cilinders', 'cilinderinhoud']
        list_milieu_keys = ['zuinigheidsclassificatie']
        list_matengewichten_keys = ['massa_ledig_voertuig','toegestane_maximum_massa_voertuig', 'massa_rijklaar', 'maximum_massa_trekken_ongeremd', 'maximum_trekken_massa_geremd', 
                                    'afstand_hart_koppeling_tot_achterzijde_voertuig', 'afstand_voorzijde_voertuig_tot_hart_koppeling', 'lengte', 'breedte', 
                                    'wielbasis', 'maximum_massa_samenstelling']

        # Het kan zijn dat bepaalde keys wel in de lists hierboven voorkomen, maar niet in de lijst met gegevens zoals die worden opgehaald bij het RDW.
        # In dat geval zijn list_registratie_keys en list_registratie_values niet even lang. 
        # Met onderstaande code verwijderen we de items die wel in list_registratie_values staan, maar die niet in de gegevens van het RDW staan. Op deze manier voorkomen we een error
        # Het is mij nog niet bekend of dit ook bij de andere categorieen voorkomen. Dit vraagt om testen. 
        for key in list_basiskenmerken_keys:
            if key not in car_property:
                # print('Niet in car_property: ' + key)
                list_basiskenmerken_keys.remove(key)
        
        for key in list_registratie_keys:
            if key not in car_property:
                # print('Niet in car_property: ' + key)
                list_registratie_keys.remove(key)

        for key in list_motor_keys:
            if key not in car_property:
                # print('Niet in car_property: ' + key)
                list_motor_keys.remove(key)

        for key in list_milieu_keys:
            if key not in car_property:
                # print('Niet in car_property: ' + key)
                list_milieu_keys.remove(key)

        for key in list_matengewichten_keys:
            if key not in car_property:
                # print('Niet in car_property: ' + key)
                list_matengewichten_keys.remove(key)

        # We gaan nu per categorie de lijst met values vullen op basis van de bijbehorende list_category_keys

        # Categorie 1: Basiskenmerken
        indices = []
        list_basiskenmerken_values = []

        for key in list_basiskenmerken_keys:
            for i in range(len(car_property)):
                if car_property[i] == key:
                    indices.append(i)
        # print(indices)

        for index in indices:
            value = car_data[index]
            list_basiskenmerken_values.append(value)
        # print(list_basiskenmerken_keys)
        # print(list_basiskenmerken_values)

        basiskenmerken = {
            'Property' : list_basiskenmerken_keys,
            'Car Data' : list_basiskenmerken_values
        }

        # Categorie 2: Registratie
        indices = []   
        list_registratie_values = []

        for key in list_registratie_keys:
            for i in range(len(car_property)):
                if car_property[i] == key:
                    indices.append(i)
        # print(indices)

        for index in indices:
            value = car_data[index]
            list_registratie_values.append(value)            
        
        registratie = {
            'Property' : list_registratie_keys,
            'Car Data' : list_registratie_values
        }
        
        # print(registratie)
        # Catgeorie 3: Motor
        indices = []
        list_motor_values = []

        for key in list_motor_keys:
            for i in range(len(car_property)):
                if car_property[i] == key:
                    indices.append(i)
        # print(indices)

        for index in indices:
            value = car_data[index]
            list_motor_values.append(value)            
        
        motor = {
            'Property' : list_motor_keys,
            'Car Data' : list_motor_values
        }
        
        # Categorie 4: Milieu
        indices = []
        list_milieu_values = []

        for key in list_milieu_keys:
            for i in range(len(car_property)):
                if car_property[i] == key:
                    indices.append(i)
        #print(indices)

        for index in indices:
            value = car_data[index]
            list_milieu_values.append(value)            
        
        milieu = {
            'Property' : list_milieu_keys,
            'Car Data' : list_milieu_values
        }

        # Categorie 5: Maten en Gewichten
        indices = []
        list_matengewichten_values = []

        for key in list_matengewichten_keys:
            for i in range(len(car_property)):
                if car_property[i] == key:
                    indices.append(i)
        # print(indices)

        for index in indices:
            value = car_data[index]
            list_matengewichten_values.append(value)            
        
        maten_en_gewichten = {
            'Property' : list_matengewichten_keys,
            'Car Data' : list_matengewichten_values
        }

        # print(car_information)
        
        # We gaan nu op basis van bovenstaande lijsten met keys de key,value paren uit het woordenboek car_information halen.
        # Deze values komen in lis_[category]_values.
        # Tot slot maken we daar dan weer een woordenboek van.    
                      
        df = pd.DataFrame(car_information)
        # print(df)
        df1 = pd.DataFrame(basiskenmerken)
        # print(df1)

        df2 = pd.DataFrame(registratie)  
        # print(df2)   

        df3 = pd.DataFrame(motor)  
        # print(df3)

        df4 = pd.DataFrame(milieu)  
        # print(df4)

        df5 = pd.DataFrame(maten_en_gewichten)  
        # print(df5)

        print("Geef nu aan welke kentekengegevens u wil zien:")
        filter0 = input("Alle categorieen? (j/n)?")
        
        if filter0 == 'j':
            print("\n")
            print("Basiskenmerken: ")
            print(df1)
            print("\n")
            print("Registratie: ")
            print(df2)
            print("\n")
            print("Motor: ")
            print(df3)
            print("\n")
            print("Milieu: ")
            print(df4)
            print("\n")
            print("Maten en Gewichten: ")
            print(df5)
        elif filter0 == "n":         
            filter1 = input("Basiskenmerken? (j/n)")
            filter2 = input("Registratie? (j/n)")
            filter3 = input("Motor? (j/n)")
            filter4 = input("Milieu? (j/n)")
            filter5 = input("Maten en Gewichten? (j/n)")

            if filter1 == "j":
                print("\n")
                print("Basiskenmerken: ")
                print(df1)
            if filter2 == "j":
                print("\n")
                print("Registratie: ")
                print(df2)
            if filter3 == "j":
                print("\n")
                print("Motor: ")
                print(df3)
            if filter4 == "j":
                print("\n")
                print("Milieu: ")
                print(df4)
            if filter5 == "j":
                print("\n")
                print("Maten en Gewichten: ")
                print(df5)


        # Initialiseer Word-document, maak lege tabel in Word, vul de tabel met de gegevens uit het DataFrame 
        # en sla het Word-kentekenrapport vervolgens op
        # We hebben nu ook een filteroptie: de gebruiker kan aangeven welke categorieen info hij in het Word-kentekenrapport wil.
        # Vandaar dat er in onderstaande code veel if-lussen voorkomen. 

        doc = docx.Document()

        heading = str('Kentekenrapport voor kenteken: ' + kenteken )
        doc.add_heading(heading, 0)
        
        if filter0 == 'j':
            doc.add_heading('Basiskenmerken', level=1)
            t1 = doc.add_table(rows = df1.shape[0], cols = df1.shape[1])
            t1.style = 'TableGrid'
        
            doc.add_heading('Registratie', level=1)
            t2 = doc.add_table(rows = df2.shape[0], cols = df2.shape[1])
            t2.style = 'TableGrid'

            doc.add_heading('Motor', level=1)
            t3 = doc.add_table(rows = df3.shape[0], cols = df3.shape[1])
            t3.style = 'TableGrid'

            doc.add_heading('Mileu', level=1)
            t4 = doc.add_table(rows = df4.shape[0], cols = df4.shape[1])
            t4.style = 'TableGrid'

            doc.add_heading('Maten en Gewichten', level=1)
            t5 = doc.add_table(rows = df5.shape[0], cols = df5.shape[1])
            t5.style = 'TableGrid'

            for i in range(df1.shape[0]):
                for j in range(df1.shape[1]):
                    cell = df1.iat[i,j]
                    t1.cell(i,j).text = str(cell)

            for i in range(df2.shape[0]):
                for j in range(df2.shape[1]):
                    cell = df2.iat[i,j]
                    t2.cell(i,j).text = str(cell)
        
            for i in range(df3.shape[0]):
                for j in range(df3.shape[1]):
                    cell = df3.iat[i,j]
                    t3.cell(i,j).text = str(cell)
        
            for i in range(df4.shape[0]):
                for j in range(df4.shape[1]):
                    cell = df4.iat[i,j]
                    t4.cell(i,j).text = str(cell)

            for i in range(df5.shape[0]):
                for j in range(df5.shape[1]):
                    cell = df5.iat[i,j]
                    t5.cell(i,j).text = str(cell)
        elif filter0 == 'n':
            if filter1 == "j":    
                doc.add_heading('Basiskenmerken', level=1)
                t1 = doc.add_table(rows = df1.shape[0], cols = df1.shape[1])
                t1.style = 'TableGrid'

                for i in range(df1.shape[0]):
                    for j in range(df1.shape[1]):
                        cell = df1.iat[i,j]
                        t1.cell(i,j).text = str(cell)               
                         
            if filter2 == "j":
                doc.add_heading('Registratie', level=1)
                t2 = doc.add_table(rows = df2.shape[0], cols = df2.shape[1])
                t2.style = 'TableGrid'
            
                for i in range(df2.shape[0]):
                    for j in range(df2.shape[1]):
                        cell = df2.iat[i,j]
                        t2.cell(i,j).text = str(cell)

            if filter3 == "j":
                doc.add_heading('Motor', level=1)
                t3 = doc.add_table(rows = df3.shape[0], cols = df3.shape[1])
                t3.style = 'TableGrid'

                for i in range(df3.shape[0]):
                    for j in range(df3.shape[1]):
                        cell = df3.iat[i,j]
                        t3.cell(i,j).text = str(cell)

            if filter4 == "j":
                doc.add_heading('Mileu', level=1)
                t4 = doc.add_table(rows = df4.shape[0], cols = df4.shape[1])
                t4.style = 'TableGrid'

                for i in range(df4.shape[0]):
                    for j in range(df4.shape[1]):
                        cell = df4.iat[i,j]
                        t4.cell(i,j).text = str(cell)

            if filter5 == "j":
                doc.add_heading('Maten en Gewichten', level=1)
                t5 = doc.add_table(rows = df5.shape[0], cols = df5.shape[1])
                t5.style = 'TableGrid'  

                for i in range(df5.shape[0]):
                    for j in range(df5.shape[1]):
                        cell = df5.iat[i,j]
                        t5.cell(i,j).text = str(cell)       

        name_file = str(kenteken + '.docx') 
        doc.save(name_file)



        
        